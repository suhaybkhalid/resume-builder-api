import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# Replace with your OpenAI API Key
OPENAI_API_KEY = "OPENAI_API_KEY"

def format_resume_with_ai(text):
    """ Uses OpenAI GPT-4 to improve resume text, enforce structure, and ensure professional formatting."""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a professional resume formatter. Format the given resume content professionally."},
            {"role": "user", "content": text}
        ],
        api_key=OPENAI_API_KEY
    )
    return response["choices"][0]["message"]["content"]

def process_resume(input_file, output_file):
    """ Reads a resume from a .docx file, formats it, and saves the improved version."""
    doc = Document(input_file)
    raw_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Process text with GPT-4 for professional formatting
    formatted_text = format_resume_with_ai(raw_text)
    
    # Create a new Word document with the improved formatting
    new_doc = Document()
    
    # Formatting the resume
    for line in formatted_text.split("\n"):
        if line.strip():
            if line.isupper():  # Section titles (bold + font size increase)
                para = new_doc.add_paragraph()
                run = para.add_run(line.strip())
                run.bold = True
                run.font.size = Pt(14)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif line.startswith("•"):  # Bullet points
                para = new_doc.add_paragraph(line.strip(), style='ListBullet')
            else:
                para = new_doc.add_paragraph(line.strip())
    
    # Save the formatted resume
    new_doc.save(output_file)
    print(f"✅ Resume formatting completed! Saved as: {output_file}")

# Run the formatter
process_resume("input_resume.docx", "Formatted_Resume.docx")

